import os
import re
import sys
import csv
import traceback
import unicodedata  # NOVO
from datetime import datetime
from tkinter import (
    Tk, Button, Label, Text, END, filedialog, Scrollbar, RIGHT, Y, LEFT, BOTH,
    DISABLED, NORMAL, messagebox
)
from pypdf import PdfReader
from docx import Document

APP_VERSION = "1.3.0"

# Regex de e-mails
REGEX_EMAIL = re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', re.IGNORECASE)

def log(txt_widget: Text, msg: str):
    """Escreve no log (área de texto), rolando automaticamente."""
    txt_widget.config(state=NORMAL)
    txt_widget.insert(END, msg + "\n")
    txt_widget.see(END)
    txt_widget.config(state=DISABLED)
    txt_widget.update_idletasks()

def extrair_emails_de_pdf(caminho_pdf: str) -> list:
    """Extrai todos os e-mails do PDF."""
    emails = []
    with open(caminho_pdf, "rb") as f:
        reader = PdfReader(f)
        texto = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            texto.append(t)
        texto = "\n".join(texto)
    emails = REGEX_EMAIL.findall(texto)
    return emails

def primeira_linha_com_cnpj(caminho_pdf: str):
    """
    Varre o PDF e retorna a PRIMEIRA linha que contém 'CNPJ' (case-sensitive),
    ou None se não encontrar.
    """
    try:
        with open(caminho_pdf, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                txt = (page.extract_text() or "")
                for linha in txt.splitlines():
                    if "CNPJ" in linha:
                        return linha.strip()
    except Exception:
        pass
    return None

def trecho_endereco_cep_quadro(caminho_pdf: str):
    """
    Procura a primeira ocorrência de 'CEP' e extrai até a próxima linha que contenha
    'QUADRO' (exclusivo), removendo 'QUADRO' se sobrar no fim.
    Retorna string ou None.
    """
    try:
        with open(caminho_pdf, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                txt = (page.extract_text() or "")
                linhas = txt.splitlines()

                cep_idx = None
                for i, linha in enumerate(linhas):
                    if "CEP" in linha:
                        cep_idx = i
                        break
                if cep_idx is None:
                    continue

                quadro_idx = None
                for j in range(cep_idx, len(linhas)):
                    if "QUADRO" in linhas[j]:
                        quadro_idx = j
                        break

                if quadro_idx is not None and quadro_idx > cep_idx:
                    bloco = linhas[cep_idx:quadro_idx]
                    # limpar 'QUADRO' caso tenha sobrado parcialmente
                    if bloco and "QUADRO" in bloco[-1]:
                        bloco[-1] = bloco[-1].replace("QUADRO", "").strip()
                    trecho = "\n".join(bloco).strip()
                    if trecho:
                        return trecho
    except Exception:
        pass
    return None

# ---------- Funções para salvar DOCX ----------

def salvar_docx_emails(emails_unicos, saida_docx):
    doc = Document()
    doc.add_heading('Lista de E-mails Encontrados', level=1)
    doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph('')
    if emails_unicos:
        for e in sorted(emails_unicos, key=str.lower):
            doc.add_paragraph(e)
    else:
        doc.add_paragraph('Nenhum e-mail encontrado.')
    doc.save(saida_docx)

def salvar_docx_cnpj(itens, saida_docx):
    """
    itens: lista de tuplas (nome_arquivo, linha_cnpj)
    Gera um DOCX SEM a linha 'Arquivo: ...', contendo apenas as linhas de CNPJ.
    """
    doc = Document()
    doc.add_heading("Linhas com CNPJ (primeira ocorrência por arquivo)", level=1)
    doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph('')

    if not itens:
        doc.add_paragraph('Nenhum dado encontrado.')
    else:
        for _, linha_cnpj in itens:
            if linha_cnpj:
                doc.add_paragraph(linha_cnpj)
            else:
                doc.add_paragraph("(vazio)")
    doc.save(saida_docx)

def salvar_docx_enderecos(itens_end, mapa_cnpj_por_arquivo, saida_docx):
    """
    itens_end: lista de tuplas (nome_arquivo, trecho_endereco_multilinha)
    mapa_cnpj_por_arquivo: dict {nome_arquivo: linha_cnpj}
    """
    doc = Document()
    doc.add_heading("Trechos entre CEP e QUADRO (primeira ocorrência por arquivo)", level=1)
    doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph('')

    if not itens_end:
        doc.add_paragraph('Nenhum dado encontrado.')
    else:
        for nome_arq, trecho in itens_end:
            linha_cnpj = mapa_cnpj_por_arquivo.get(nome_arq)
            if linha_cnpj:
                doc.add_paragraph(linha_cnpj)
            if trecho:
                for linha in str(trecho).splitlines():
                    doc.add_paragraph(linha)
            else:
                doc.add_paragraph("(vazio)")
            doc.add_paragraph("")
    doc.save(saida_docx)

# ---------- NOVO: utilitários de normalização/pesquisa ----------

def _normalize(s: str) -> str:
    """Remove acentos e coloca em minúsculas para busca robusta."""
    if s is None:
        return ""
    nfkd = unicodedata.normalize("NFKD", s)
    s_sem_acento = "".join(ch for ch in nfkd if not unicodedata.combining(ch))
    return s_sem_acento.lower()

# ---------- NOVO: extração Razão Social e CNPJ até 'R' ----------

def extrair_razao_social(caminho_pdf: str):
    """
    Procura a 1ª linha que contém 'Razão Social' (variações com/sem acento e caixa)
    e retorna APENAS o texto à direita do marcador na MESMA linha.
    """
    variantes = ["Razão Social", "RAZÃO SOCIAL", "Razao Social", "RAZAO SOCIAL"]
    try:
        with open(caminho_pdf, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                txt = (page.extract_text() or "")
                for linha in txt.splitlines():
                    for v in variantes:
                        if v in linha:
                            pos = linha.find(v) + len(v)
                            # Ignora pontuação comum entre label e valor
                            valor = linha[pos:].lstrip(" :\t-").strip()
                            if valor:
                                return valor
                    # fallback: busca normalizada (sem acento / minúscula)
                    if "razao social" in _normalize(linha):
                        # tentar achar índice aproximado na string original
                        # usando 'Razao Social' (sem acento) como referência
                        idx_norm = _normalize(linha).find("razao social")
                        # recuperar comprimento da label original "Razão Social"
                        pos_aprox = idx_norm + len("razao social")
                        # como é aproximado, apenas corta uma quantidade segura:
                        valor = linha[pos_aprox:].lstrip(" :\t-").strip()
                        if valor:
                            return valor
    except Exception:
        pass
    return None

def extrair_cnpj_ate_R(caminho_pdf: str):
    """
    Procura a 1ª ocorrência de 'CNPJ' e retorna a substring
    DE 'CNPJ' ATÉ a primeira letra 'R' após ele (SEM incluir 'R').
    Caso não encontre 'R' após 'CNPJ', retorna da ocorrência até o fim da linha.
    """
    try:
        with open(caminho_pdf, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                txt = (page.extract_text() or "")
                for linha in txt.splitlines():
                    if "CNPJ" in linha:
                        start = linha.find("CNPJ")
                        sub = linha[start:]
                        # achar a primeira letra 'R' após 'CNPJ'
                        idx_r = sub.find("R")
                        if idx_r > 0:
                            trecho = sub[:idx_r].rstrip()
                        else:
                            trecho = sub.strip()
                        # opcional: remover espaços/pontuação supérflua no fim
                        return trecho.rstrip(" -:;")
    except Exception:
        pass
    return None

# ---------- NOVO: salvar DOCX (Razão + CNPJ) ----------

def salvar_docx_razao_cnpj(pares, saida_docx):
    """
    pares: lista de tuplas (razao_social, cnpj_trecho)
    Grava uma linha por arquivo no formato:
      <Razão Social> <TAB> <CNPJ...até 'R'>
    """
    doc = Document()
    doc.add_heading("Razão Social e CNPJ (1ª ocorrência por arquivo)", level=1)
    doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph('')

    if not pares:
        doc.add_paragraph('Nenhum dado encontrado.')
    else:
        for razao, cnpj_trecho in pares:
            rz = razao if razao else "(Razão Social não encontrada)"
            cj = cnpj_trecho if cnpj_trecho else "(CNPJ não encontrado)"
            # mesma linha, separados por TAB
            doc.add_paragraph(f"{rz}\t{cj}")

    doc.save(saida_docx)

# ---------------------------------------------

def processar(btn_processar, txt_log):
    try:
        btn_processar.config(state=DISABLED)

        # limpar log
        txt_log.config(state=NORMAL)
        txt_log.delete(1.0, END)
        txt_log.config(state=DISABLED)

        # selecionar PDFs
        arquivos = filedialog.askopenfilenames(
            title="Selecione um ou mais PDFs",
            filetypes=[("Arquivos PDF", "*.pdf")]
        )
        if not arquivos:
            log(txt_log, "⚠️ Nenhum PDF selecionado.")
            return

        log(txt_log, f"📄 PDFs selecionados: {len(arquivos)}")

        # e-mails (lista geral)
        todos_emails = []

        # coletores
        cnpj_itens = []   # (arquivo, linha CNPJ)
        end_itens  = []   # (arquivo, trecho CEP..QUADRO)

        # loop principal
        for idx, caminho in enumerate(arquivos, start=1):
            nome_arq = os.path.basename(caminho)
            log(txt_log, f"• ({idx}/{len(arquivos)}) {nome_arq}")

            # e-mails
            try:
                emails = extrair_emails_de_pdf(caminho)
                log(txt_log, f"   → {len(emails)} e-mail(s)")
                todos_emails.extend(emails)
            except Exception as e:
                log(txt_log, f"   ✗ Erro ao extrair e-mails: {e}")

            # CNPJ (linha inteira - modo antigo)
            try:
                linha_cnpj = primeira_linha_com_cnpj(caminho)
                if linha_cnpj:
                    cnpj_itens.append((nome_arq, linha_cnpj))
                    log(txt_log, "   → CNPJ: OK")
                else:
                    log(txt_log, "   → CNPJ: não encontrado")
            except Exception as e:
                log(txt_log, f"   ✗ Erro no CNPJ: {e}")

            # Endereço CEP..QUADRO
            try:
                trecho = trecho_endereco_cep_quadro(caminho)
                if trecho:
                    end_itens.append((nome_arq, trecho))
                    log(txt_log, "   → Endereço (CEP..QUADRO): OK")
                else:
                    log(txt_log, "   → Endereço: não encontrado")
            except Exception as e:
                log(txt_log, f"   ✗ Erro no endereço: {e}")

        # remove duplicados de e-mail
        unicos = sorted(set(todos_emails), key=str.lower)
        log(txt_log, f"\n📬 Total extraídos: {len(todos_emails)}  |  Únicos: {len(unicos)}")

        # pasta de saída = pasta do primeiro PDF selecionado
        pasta_saida = os.path.dirname(arquivos[0]) if arquivos else os.getcwd()

        # ---------- MAPA: nome_arquivo -> linha CNPJ (para uso nos endereços)
        cnpj_por_arquivo = {nome: linha for (nome, linha) in cnpj_itens}

        # --------- Salvar 3 DOCX ---------
        # 1) E-mails
        saida_emails_docx = os.path.join(pasta_saida, "emails_encontrados.docx")
        try:
            salvar_docx_emails(unicos, saida_emails_docx)
        except Exception:
            log(txt_log, "⚠️ Não foi possível gravar o DOCX de e-mails nessa pasta. Escolha outro local…")
            alt = filedialog.asksaveasfilename(
                title="Salvar DOCX de e-mails como",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialfile="emails_encontrados.docx"
            )
            if alt:
                salvar_docx_emails(unicos, alt)
                saida_emails_docx = alt
            else:
                saida_emails_docx = None

        # 2) CNPJ (sem 'Arquivo: ...')
        saida_cnpj_docx = os.path.join(pasta_saida, "extrair_cnpj_nome.docx")
        try:
            salvar_docx_cnpj(cnpj_itens, saida_cnpj_docx)
        except Exception:
            log(txt_log, "⚠️ Não foi possível gravar o DOCX de CNPJ nessa pasta. Escolha outro local…")
            alt = filedialog.asksaveasfilename(
                title="Salvar DOCX (CNPJ) como",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialfile="extrair_cnpj_nome.docx"
            )
            if alt:
                salvar_docx_cnpj(cnpj_itens, alt)
                saida_cnpj_docx = alt
            else:
                saida_cnpj_docx = None

        # 3) Endereços (sem 'Arquivo: ...' e com a 1ª linha CNPJ no topo)
        saida_end_docx = os.path.join(pasta_saida, "extrair_enderecos.docx")
        try:
            salvar_docx_enderecos(end_itens, cnpj_por_arquivo, saida_end_docx)
        except Exception:
            log(txt_log, "⚠️ Não foi possível gravar o DOCX de endereços nessa pasta. Escolha outro local…")
            alt = filedialog.asksaveasfilename(
                title="Salvar DOCX (Endereços) como",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialfile="extrair_enderecos.docx"
            )
            if alt:
                salvar_docx_enderecos(end_itens, cnpj_por_arquivo, alt)
                saida_end_docx = alt
            else:
                saida_end_docx = None

        # mensagens finais
        msg_ok = "Processo concluído."
        if saida_emails_docx: msg_ok += f"\nDOCX (e-mails): {saida_emails_docx}"
        if saida_cnpj_docx:   msg_ok += f"\nDOCX (CNPJ):   {saida_cnpj_docx}"
        if saida_end_docx:    msg_ok += f"\nDOCX (End.):   {saida_end_docx}"
        messagebox.showinfo("Concluído", msg_ok)

    except Exception as e:
        log(txt_log, "❌ Falha inesperada. Detalhes no traceback abaixo:")
        log(txt_log, traceback.format_exc())
        messagebox.showerror("Erro", str(e))
    finally:
        btn_processar.config(state=NORMAL)

# ---------- NOVO: fluxo do 2º botão (Razão + CNPJ) ----------

def processar_razao_cnpj(btn, txt_log):
    try:
        btn.config(state=DISABLED)

        # limpar log
        txt_log.config(state=NORMAL)
        txt_log.delete(1.0, END)
        txt_log.config(state=DISABLED)

        # selecionar PDFs
        arquivos = filedialog.askopenfilenames(
            title="Selecione um ou mais PDFs",
            filetypes=[("Arquivos PDF", "*.pdf")]
        )
        if not arquivos:
            log(txt_log, "⚠️ Nenhum PDF selecionado.")
            return

        log(txt_log, f"📄 PDFs selecionados: {len(arquivos)}")

        pares = []  # (razao_social, cnpj_trecho)

        for idx, caminho in enumerate(arquivos, start=1):
            nome_arq = os.path.basename(caminho)
            log(txt_log, f"• ({idx}/{len(arquivos)}) {nome_arq}")

            try:
                razao = extrair_razao_social(caminho)
                log(txt_log, f"   → Razão Social: {'OK' if razao else 'não encontrada'}")
            except Exception as e:
                log(txt_log, f"   ✗ Erro na Razão Social: {e}")
                razao = None

            try:
                cnpj = extrair_cnpj_ate_R(caminho)
                log(txt_log, f"   → CNPJ (até 'R'): {'OK' if cnpj else 'não encontrado'}")
            except Exception as e:
                log(txt_log, f"   ✗ Erro no CNPJ (até 'R'): {e}")
                cnpj = None

            pares.append((razao, cnpj))

        # pasta de saída = pasta do primeiro PDF selecionado
        pasta_saida = os.path.dirname(arquivos[0]) if arquivos else os.getcwd()
        saida_docx = os.path.join(pasta_saida, "extrair_cnpj_nome.docx")

        try:
            salvar_docx_razao_cnpj(pares, saida_docx)
            messagebox.showinfo("Concluído", f"Processo concluído.\nDOCX: {saida_docx}")
        except Exception as e:
            log(txt_log, "⚠️ Não foi possível gravar o DOCX nessa pasta. Escolha outro local…")
            alt = filedialog.asksaveasfilename(
                title="Salvar DOCX (Razão + CNPJ) como",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialfile="extrair_cnpj_nome.docx"
            )
            if alt:
                salvar_docx_razao_cnpj(pares, alt)
                messagebox.showinfo("Concluído", f"Processo concluído.\nDOCX: {alt}")
            else:
                messagebox.showwarning("Aviso", "Arquivo não salvo.")
    except Exception:
        log(txt_log, "❌ Falha inesperada. Detalhes no traceback abaixo:")
        log(txt_log, traceback.format_exc())
        messagebox.showerror("Erro", "Falha inesperada (ver log).")
    finally:
        btn.config(state=NORMAL)

def main():
    root = Tk()
    root.title(f"Extrator de PDFs — v{APP_VERSION}")
    root.geometry("760x560")

    Label(root, text="Selecione PDFs e clique em Processar para gerar 3 DOCX: e-mails, CNPJ e endereços.").pack(pady=8)

    # botão principal (modo original)
    btn_processar = Button(root, text="Selecionar PDFs e Processar", width=32,
                           command=lambda: processar(btn_processar, txt_log))
    btn_processar.pack(pady=6)

    # ---------- ALTERAÇÃO NO main(): novo botão ----------
    btn_razao = Button(root, text="Botão do Valverde", width=32)
    btn_razao.config(command=lambda: processar_razao_cnpj(btn_razao, txt_log))
    btn_razao.pack(pady=6)

    # log (mantemos o painel, sem mensagem inicial)
    Label(root, text="Log:").pack(anchor="w", padx=8)
    txt_log = Text(root, height=20, state=DISABLED, wrap="word")
    scroll = Scrollbar(root, command=txt_log.yview)
    txt_log.configure(yscrollcommand=scroll.set)
    txt_log.pack(side=LEFT, fill=BOTH, expand=True, padx=(8,0), pady=(0,8))
    scroll.pack(side=RIGHT, fill=Y, pady=(0,8))

    root.mainloop()

if __name__ == "__main__":
    # Ajuste de diretório quando empacotado (PyInstaller)
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        os.chdir(os.path.dirname(sys.executable))
    main()
